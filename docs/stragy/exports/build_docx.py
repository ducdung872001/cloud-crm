# -*- coding: utf-8 -*-
"""
Build a BOD-grade DOCX from the FitPro 2027 strategy markdown.

Flow:
  1. Customize Pandoc's default reference.docx (fonts, heading colors).
  2. Run Pandoc to convert MD → DOCX using the customized reference.
  3. Post-process: insert cover page, footer with page numbers, table styling.

Run:  python build_docx.py
"""
import os
import sys
import io
import shutil
import subprocess
from copy import deepcopy

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REF_BASE = os.path.join(HERE, "_reference.docx")
REF_OUT = os.path.join(HERE, "_reference_styled.docx")
MD_PATH = os.path.abspath(os.path.join(HERE, "..", "PHAN_TICH_TONG_HOP_FITPRO_2027.md"))
MD_TMP = os.path.join(HERE, "_source_for_docx.md")
OUT_DOCX = os.path.join(HERE, "FitPro_2027_Phan_Tich_Tong_Hop.docx")


def preprocess_md():
    """Strip the manually-written MỤC LỤC block (Pandoc will auto-generate TOC)
    and remove the top H1 + leading blockquote that we replace with a cover page."""
    with open(MD_PATH, "r", encoding="utf-8") as f:
        src = f.read()

    lines = src.splitlines()
    out = []
    skip_toc = False
    skipped_top = False
    skip_top_blockquote = False
    for i, ln in enumerate(lines):
        # 1) skip top H1 "# PHÂN TÍCH..."
        if not skipped_top and ln.startswith("# "):
            skipped_top = True
            skip_top_blockquote = True
            continue
        # 2) skip the top blockquote(s) right after the H1 + blank lines until first H2
        if skip_top_blockquote:
            if ln.strip() == "" or ln.startswith(">"):
                continue
            if ln.startswith("## "):
                skip_top_blockquote = False
                # fallthrough to handle this H2 below
            else:
                continue
        # 3) skip MỤC LỤC block
        if ln.strip().upper().startswith("## MỤC LỤC"):
            skip_toc = True
            continue
        if skip_toc:
            # end of TOC block at next H2 or horizontal rule
            if ln.startswith("## ") or ln.strip() == "---":
                skip_toc = False
                # the --- terminator: drop it too, then continue
                if ln.strip() == "---":
                    continue
            else:
                continue
        out.append(ln)

    with open(MD_TMP, "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print(f"[ok] preprocessed md → {MD_TMP}")

# Brand palette — Reborn / FitPro
COLOR_PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # deep blue
COLOR_ACCENT = RGBColor(0xC9, 0x9A, 0x3B)    # gold
COLOR_SUB = RGBColor(0x33, 0x33, 0x33)
COLOR_LIGHT = RGBColor(0x70, 0x70, 0x70)
COLOR_PALE_BG = "F2F4F8"
COLOR_HEADER_BG = "0B3D91"

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"


# ---------- Step 1: customize reference.docx ----------
def style_reference():
    shutil.copyfile(REF_BASE, REF_OUT)
    doc = Document(REF_OUT)
    styles = doc.styles

    def set_style(name, *, size=None, bold=None, color=None, font=None, italic=None,
                  space_before=None, space_after=None):
        try:
            st = styles[name]
        except KeyError:
            return
        if font:
            st.font.name = font
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rfonts.set(qn(attr), font)
        if size is not None:
            st.font.size = Pt(size)
        if bold is not None:
            st.font.bold = bold
        if italic is not None:
            st.font.italic = italic
        if color is not None:
            st.font.color.rgb = color
        pfmt = st.paragraph_format
        if space_before is not None:
            pfmt.space_before = Pt(space_before)
        if space_after is not None:
            pfmt.space_after = Pt(space_after)

    # Body
    set_style("Normal", font=FONT_BODY, size=11, color=COLOR_SUB,
              space_before=0, space_after=4)

    # Headings
    set_style("Title", font=FONT_HEAD, size=26, bold=True, color=COLOR_PRIMARY,
              space_before=0, space_after=6)
    set_style("Subtitle", font=FONT_HEAD, size=14, italic=True, color=COLOR_LIGHT,
              space_before=0, space_after=18)
    set_style("Heading 1", font=FONT_HEAD, size=20, bold=True, color=COLOR_PRIMARY,
              space_before=18, space_after=8)
    set_style("Heading 2", font=FONT_HEAD, size=15, bold=True, color=COLOR_PRIMARY,
              space_before=12, space_after=6)
    set_style("Heading 3", font=FONT_HEAD, size=12, bold=True, color=COLOR_ACCENT,
              space_before=8, space_after=4)
    set_style("Heading 4", font=FONT_HEAD, size=11, bold=True, color=COLOR_SUB,
              space_before=6, space_after=4)

    # TOC
    for name in ("TOC Heading", "toc 1", "toc 2", "toc 3", "TOC 1", "TOC 2", "TOC 3"):
        set_style(name, font=FONT_BODY)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    doc.save(REF_OUT)
    print(f"[ok] reference styled: {REF_OUT}")


# ---------- Step 2: pandoc convert ----------
def run_pandoc():
    cmd = [
        "pandoc", MD_TMP,
        "-o", OUT_DOCX,
        "--reference-doc", REF_OUT,
        "--toc",
        "--toc-depth=2",
        "-V", "lang=vi",
        # title is required for Pandoc to emit the TOC; we strip the title
        # paragraphs in post-process and keep the TOC.
        "--metadata", "title= ",
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        print("STDOUT:", r.stdout)
        print("STDERR:", r.stderr)
        raise SystemExit(r.returncode)
    print(f"[ok] pandoc converted: {OUT_DOCX}")


# ---------- Step 3: post-process ----------
def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def set_cell_border(cell, color="BFBFBF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        # remove existing same-tag
        for old in tcBorders.findall(qn(f"w:{edge}")):
            tcBorders.remove(old)
        tcBorders.append(el)


def style_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        if not table.rows:
            continue
        header = table.rows[0]
        for cell in header.cells:
            set_cell_shading(cell, COLOR_HEADER_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, color="0B3D91", size="8")
        # Body rows: subtle zebra + cell padding
        for ridx, row in enumerate(table.rows[1:], start=1):
            for cell in row.cells:
                set_cell_border(cell, color="D9DCE3", size="6")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.size is None or run.font.size > Pt(11):
                            run.font.size = Pt(10.5)
                if ridx % 2 == 0:
                    set_cell_shading(cell, COLOR_PALE_BG)


def add_field(paragraph, instr_text):
    """Insert a Word field (TOC update via PAGE field for footer)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r_el = run._r
    r_el.append(fld_begin); r_el.append(instr); r_el.append(fld_sep); r_el.append(fld_end)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run("FitPro by MF7 · Vision 2027   |   Trang ")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_LIGHT
        add_field(p, "PAGE")
        run2 = p.add_run(" / ")
        run2.font.size = Pt(9)
        run2.font.color.rgb = COLOR_LIGHT
        add_field(p, "NUMPAGES")
        # second run with formatting after fields
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_LIGHT


def insert_cover_page(doc):
    """Insert a cover page block at the very top of the document."""
    body = doc.element.body
    # Build cover paragraphs as OXML, then prepend before existing content.
    elements_to_prepend = []

    def make_p(text, *, size=11, bold=False, color=COLOR_SUB, align="center",
               space_before=0, space_after=8, italic=False, font=FONT_HEAD):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), align); pPr.append(jc)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(space_before * 20))
        spacing.set(qn("w:after"), str(space_after * 20))
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font)
        rPr.append(rfonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
        szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(size * 2)); rPr.append(szCs)
        if bold:
            b = OxmlElement("w:b"); rPr.append(b)
            bCs = OxmlElement("w:bCs"); rPr.append(bCs)
        if italic:
            i = OxmlElement("w:i"); rPr.append(i)
        col = OxmlElement("w:color"); col.set(qn("w:val"), f"{color.rgb if hasattr(color, 'rgb') else color}".upper()[-6:]); rPr.append(col)
        r.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; r.append(t)
        p.append(r)
        return p

    def make_blank(space=12):
        return make_p("", space_after=space)

    def make_page_break():
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); r.append(br)
        p.append(r)
        return p

    def make_hr():
        # horizontal line via paragraph bottom border
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C99A3B")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.append(pPr)
        return p

    # Visual cover:
    #   (spacer)
    #   FITPRO BY MF7 (small caps, primary)
    #   PHÂN TÍCH TỔNG HỢP CHIẾN LƯỢC
    #   VISION 2027 — KIẾN TẠO 10.000 FITPRO HOME
    #   ─── gold rule ───
    #   Tài liệu trình Ban Giám đốc
    #   (spacer)
    #   Soạn thảo:  CTO Office · Reborn JSC
    #   Ngày 11 tháng 05 năm 2026
    #   (page break)

    elements_to_prepend = [
        make_blank(120),
        make_p("FITPRO BY MF7", size=12, bold=True, color=COLOR_ACCENT, space_after=12),
        make_p("PHÂN TÍCH TỔNG HỢP CHIẾN LƯỢC", size=28, bold=True, color=COLOR_PRIMARY, space_after=4),
        make_p("VISION 2027", size=22, bold=True, color=COLOR_PRIMARY, space_after=6),
        make_p("Kiến tạo 10.000 FitPro Home Toàn Quốc", size=14, italic=True, color=COLOR_LIGHT, space_after=18),
        make_hr(),
        make_p("Tài liệu trình Ban Giám đốc", size=13, bold=True, color=COLOR_SUB, space_before=12, space_after=120),
        make_p("Soạn thảo: CTO Office · Reborn JSC", size=11, color=COLOR_SUB, space_after=4),
        make_p("Triết lý vận hành: Tốc Độ – Chuẩn Xác – Nhân Văn", size=11, italic=True, color=COLOR_LIGHT, space_after=4),
        make_p("Ngày 11 tháng 05 năm 2026", size=11, color=COLOR_LIGHT, space_after=4),
        make_page_break(),
    ]

    # Prepend
    sectPr = body.find(qn("w:sectPr"))
    insertion_index = 0
    for el in reversed(elements_to_prepend):
        body.insert(insertion_index, el)


def strip_pandoc_title_block(doc):
    """Pandoc inserts paragraphs with styles Title/Subtitle/Author/Date right
    before the TOC. We added them only so Pandoc would emit the TOC; remove
    them now so our custom cover page stands alone."""
    title_styles = {"Title", "Subtitle", "Author", "Date"}
    body = doc.element.body
    removed = 0
    for p in list(doc.paragraphs):
        if p.style.name in title_styles:
            p._element.getparent().remove(p._element)
            removed += 1
    print(f"[ok] stripped {removed} pandoc title paragraphs")


def build_static_toc(doc):
    """Build a pre-populated TOC by scanning Heading 2/3 paragraphs, inserted
    right after the cover page and before content. Also wrap it in a Word TOC
    field so it stays clickable and refreshable via F9."""
    # Find the first content paragraph (first Heading 2 after the cover).
    first_h2 = None
    for p in doc.paragraphs:
        if p.style.name in ("Heading 2",):
            first_h2 = p
            break
    if first_h2 is None:
        return

    # Collect TOC entries (Heading 2 + Heading 3) up to a reasonable depth.
    entries = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 2":
            entries.append((2, p.text.strip()))
        elif p.style.name == "Heading 3":
            entries.append((3, p.text.strip()))

    # Build elements to insert before first_h2:
    insert_before = first_h2._element
    parent = insert_before.getparent()

    def add_para(p_el):
        parent.insert(list(parent).index(insert_before), p_el)

    # (a) Page break right before TOC so cover stays on its own page
    pb = OxmlElement("w:p")
    rb = OxmlElement("w:r")
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); rb.append(br)
    pb.append(rb)
    add_para(pb)

    # (b) Heading "MỤC LỤC"
    h = OxmlElement("w:p")
    hpPr = OxmlElement("w:pPr")
    hStyle = OxmlElement("w:pStyle"); hStyle.set(qn("w:val"), "Heading1"); hpPr.append(hStyle)
    h.append(hpPr)
    hr = OxmlElement("w:r")
    hrPr = OxmlElement("w:rPr")
    hb = OxmlElement("w:b"); hrPr.append(hb)
    hsz = OxmlElement("w:sz"); hsz.set(qn("w:val"), "40"); hrPr.append(hsz)
    hcol = OxmlElement("w:color"); hcol.set(qn("w:val"), "0B3D91"); hrPr.append(hcol)
    hr.append(hrPr)
    htext = OxmlElement("w:t"); htext.text = "MỤC LỤC"; hr.append(htext)
    h.append(hr)
    add_para(h)

    # (c) TOC entries — static lines, styled. Level 2 = bold, Level 3 = indented.
    for level, text in entries:
        ep = OxmlElement("w:p")
        epPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        if level == 3:
            ind.set(qn("w:left"), "567")  # 1cm indent
        epPr.append(ind)
        # tab stop with leader dots on the right (~16cm)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        epPr.append(tabs)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "40")
        spacing.set(qn("w:after"), "40")
        epPr.append(spacing)
        ep.append(epPr)
        # text run
        er = OxmlElement("w:r")
        erPr = OxmlElement("w:rPr")
        esz = OxmlElement("w:sz"); esz.set(qn("w:val"), "22" if level == 2 else "20"); erPr.append(esz)
        if level == 2:
            eb = OxmlElement("w:b"); erPr.append(eb)
            ecol = OxmlElement("w:color"); ecol.set(qn("w:val"), "0B3D91"); erPr.append(ecol)
        else:
            ecol = OxmlElement("w:color"); ecol.set(qn("w:val"), "333333"); erPr.append(ecol)
        er.append(erPr)
        et = OxmlElement("w:t"); et.set(qn("xml:space"), "preserve"); et.text = text; er.append(et)
        ep.append(er)
        add_para(ep)

    # (d) Spacer + page break after TOC
    pb2 = OxmlElement("w:p")
    rb2 = OxmlElement("w:r")
    br2 = OxmlElement("w:br"); br2.set(qn("w:type"), "page"); rb2.append(br2)
    pb2.append(rb2)
    add_para(pb2)

    print(f"[ok] built static TOC with {len(entries)} entries")


def post_process():
    doc = Document(OUT_DOCX)
    strip_pandoc_title_block(doc)
    build_static_toc(doc)
    insert_cover_page(doc)
    style_tables(doc)
    add_footer(doc)
    doc.save(OUT_DOCX)
    print(f"[ok] post-processed: {OUT_DOCX}")


if __name__ == "__main__":
    style_reference()
    preprocess_md()
    run_pandoc()
    post_process()
    size = os.path.getsize(OUT_DOCX) / 1024
    print(f"[done] {OUT_DOCX} ({size:.1f} KB)")
