"""Generate retest guide docx for tester."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_heading("Hướng dẫn retest — 9 bug đã fix độc lập FE", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run("Commit: ")
r.bold = True
p.add_run("81b76120 (build: fix 7 bug retest) + các commit trước\n")
r = p.add_run("Branch: ")
r.bold = True
p.add_run("reborn-retail\n")
r = p.add_run("Ngày: ")
r.bold = True
p.add_run("2026-04-15")

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "Các bug dưới đây đã fix ở FE và KHÔNG phụ thuộc backend. "
    "Tester có thể test ngay sau khi FE deploy. "
    "Những bug cần chờ BE được liệt kê riêng ở mục cuối."
).italic = True

# ── Bugs có thể test ngay ───────────────────────────────────────────────────
doc.add_heading("📋 9 bug sẵn sàng retest", level=1)

bugs = [
    {
        "ma": "C.5.3",
        "title": "Shipping: báo lỗi rõ ràng + chặn 500",
        "root_cause": (
            "BE trả HTTP 500 khi carrierCode rỗng, weight=0 hoặc các field "
            "bắt buộc invalid. FE hiển thị raw error gây khó hiểu."
        ),
        "fix": (
            "Pre-validate ở FE: chặn trước khi gọi API nếu chưa chọn hãng "
            "vận chuyển hoặc weight ≤ 0. Catch HTTP 500 từ BE và dịch sang "
            "message thân thiện tiếng Việt (6 loại lỗi thường gặp: COD vượt "
            "mức, không hỗ trợ địa điểm, hóa đơn đã vận chuyển, weight, "
            "token, network)."
        ),
        "steps": [
            "Vào Bán hàng & Đơn hàng → Giao hàng & Vận chuyển → Tạo đơn vận chuyển",
            "Case 1 — Không chọn hãng VC, bấm 'Đẩy qua Hãng Vận Chuyển' → "
            "FE phải báo 'Hãng vận chuyển chưa được hỗ trợ hoặc chưa chọn' "
            "(KHÔNG còn 500)",
            "Case 2 — Chọn hãng + weight = 0 → FE báo 'Vui lòng nhập trọng "
            "lượng hợp lệ (> 0 gram)'",
            "Case 3 — Đủ field → nếu BE trả lỗi 'COD vượt' → FE hiển thị "
            "'Giá trị COD vượt quá mức tối đa mà hãng vận chuyển cho phép "
            "(COD hiện tại: X₫). Vui lòng giảm COD hoặc chọn hãng khác.'",
            "Case 4 — BE 500 không xác định → FE hiển thị 'Hãng vận chuyển "
            "đang trả lỗi hệ thống. Vui lòng kiểm tra lại địa chỉ / khối "
            "lượng / COD và thử lại.'",
        ],
        "file": "src/pages/ShipingManagement/AddShippingOrder/AddShippingOrder.tsx",
    },
    {
        "ma": "B.1",
        "title": "Click thông báo 'Cảnh báo ngưỡng tồn kho' → 404",
        "root_cause": (
            "Handler notification chưa có case cho type INVENTORY_THRESHOLD_ALERT, "
            "dẫn đến click rơi vào default → navigate sai route → 404."
        ),
        "fix": (
            "Thêm case INVENTORY_THRESHOLD_ALERT trong handleNotificationClick "
            "(cả bell dropdown header và trang /notification): navigate tới "
            "/product_inventory và truyền highlightProductId nếu có."
        ),
        "steps": [
            "Tạo một sản phẩm có variant chạm ngưỡng tồn kho (để BE bắn "
            "notification INVENTORY_THRESHOLD_ALERT)",
            "Click vào chuông thông báo ở header → chọn thông báo 'Cảnh báo "
            "ngưỡng tồn kho' → phải chuyển sang màn 'Sản phẩm tồn kho' "
            "(/product_inventory), KHÔNG còn 404",
            "Thử lại từ trang /notification (danh sách thông báo đầy đủ) → "
            "cùng kết quả",
        ],
        "file": "src/components/header/header.tsx, src/pages/NotificationList/NotificationList.tsx",
    },
    {
        "ma": "E.1.1",
        "title": "Customer search theo SĐT/email",
        "root_cause": (
            "BE filter customer chỉ hỗ trợ search theo tên (keyword), không "
            "hỗ trợ phone và email."
        ),
        "fix": (
            "Rewrite CustomerService.filter/listshared: detect pattern keyword "
            "— chứa '@' → email, toàn số ≥6 ký tự → phone. Khi detect, fetch "
            "rộng (limit 200, keyword rỗng) rồi filter CLIENT-SIDE theo các "
            "field email/customerEmail/emailAddress/phone/number_phone/"
            "phoneNumber/customerPhone. Không phụ thuộc BE."
        ),
        "steps": [
            "Vào Khách hàng & Đối tác → Khách hàng",
            "Case 1 — Search bằng email đầy đủ, vd: 'test@reborn.vn' → phải "
            "tìm thấy KH có email đó",
            "Case 2 — Search bằng email một phần, vd: '@gmail' → danh sách "
            "filter các KH có email chứa '@gmail'",
            "Case 3 — Search bằng số điện thoại đầy đủ, vd: '0901234567' → "
            "phải tìm thấy",
            "Case 4 — Search tên bình thường, vd: 'Hoa' → vẫn hoạt động như cũ",
        ],
        "file": "src/services/CustomerService.ts",
    },
    {
        "ma": "F.6.2",
        "title": "QR Thu nợ — chia sẻ qua Zalo/FB/Messenger",
        "root_cause": (
            "Nút 'Sao chép link' trước đây copy text mô tả (không phải link), "
            "các nút Zalo/Messenger/FB mở share URL chỉ với text, không kèm "
            "ảnh QR."
        ),
        "fix": (
            "Nút 'Sao chép ảnh QR' dùng ClipboardItem để copy BLOB ảnh QR "
            "vào clipboard → dán thẳng vào Zalo/Messenger. Các nút Zalo/"
            "Messenger/FB ưu tiên Web Share API với file đính kèm (mobile), "
            "fallback: download ảnh + copy text hướng dẫn (desktop)."
        ),
        "steps": [
            "Vào Tài chính & Thanh toán → Quản lý công nợ → chọn 1 KH có "
            "công nợ → 'QR Thu nợ'",
            "Case 1 (Desktop) — Bấm 'Sao chép ảnh QR' → paste Ctrl+V vào "
            "Zalo/Messenger desktop → ảnh QR dán trực tiếp, KHÔNG còn là text",
            "Case 2 (Mobile) — Bấm 'Zalo' / 'Messenger' → mở native share "
            "sheet với ảnh QR đính kèm, chọn contact để gửi",
            "Case 3 (Desktop) — Bấm 'Zalo' → tự động tải ảnh QR về máy + "
            "copy nội dung text + mở Zalo Web → user đính ảnh đã tải",
            "Case 4 — Bấm 'Tải ảnh QR' → ảnh tải về máy với tên 'qr-thu-no-*.png'",
        ],
        "file": "src/pages/Finance/DebtManagement/index.tsx",
    },
    {
        "ma": "D.1.4",
        "title": "Nhân bản sản phẩm giữ đủ biến thể + attribute",
        "root_cause": (
            "handleDuplicate gửi optionValueIds của sản phẩm GỐC (ID thuộc về "
            "product cũ), khi BE tạo product mới với options mới sẽ không "
            "khớp ID → variant mất attribute. Ngoài ra thiếu field trackStock/"
            "stock so với flow save thường."
        ),
        "fix": (
            "Bỏ optionValueIds trong dupVariants (BE sẽ tự tạo ID mới và "
            "match theo attributes name+value). Bổ sung trackStock và stock "
            "vào body duplicate để đồng nhất với save thường."
        ),
        "steps": [
            "Vào Cài đặt bán hàng → Danh sách sản phẩm",
            "Chọn 1 SP có biến thể (vd: Màu: Đỏ/Xanh × Size: S/M/L) có "
            "barcode, giá, ảnh đầy đủ",
            "Bấm 'Nhân bản' → chờ toast 'Nhân bản thành công'",
            "Mở sản phẩm mới (có hậu tố '(Copy)') → vào Chi tiết",
            "Verify: danh sách biến thể hiện đầy đủ tất cả combination, "
            "mỗi biến thể có barcode mới (khác gốc), ảnh/giá/đơn vị giữ "
            "nguyên, attribute name+value đúng",
            "Edit 1 biến thể (đổi giá) → Lưu → mở lại → giá mới giữ nguyên "
            "(không revert)",
        ],
        "file": "src/pages/SettingSell/partials/Product/partials/AddProductPage.tsx",
    },
    {
        "ma": "D.1.5",
        "title": "Thuế suất bán không cộng thêm vào tổng hóa đơn",
        "root_cause": (
            "POS đang cộng thuế vào tổng (subtotal + tax), trong khi giá bán "
            "lẻ VN đã bao gồm VAT. Kết quả: hóa đơn cao hơn giá niêm yết."
        ),
        "fix": (
            "VAT bóc tách từ giá bán (công thức vat = gross * rate / (100 + "
            "rate)) chỉ để ghi nhận/hiển thị, KHÔNG cộng vào total. Áp dụng "
            "cho Cart, PayModal, ReceiptModal."
        ),
        "steps": [
            "Vào Bán hàng tại quầy",
            "Thêm 1 SP có thuế suất 10% (vd: Chè Thái Nguyên 100,000₫)",
            "Verify: Tổng thanh toán = 100,000₫ (không phải 110,000₫)",
            "Trong biên lai, mục 'Thuế suất' hiển thị thông tin VAT bóc "
            "tách ~9,091₫ (chỉ ghi chú, không cộng)",
            "Test với nhiều SP khác thuế suất → tổng vẫn = sum(price×qty)",
        ],
        "file": "src/pages/CounterSales/components/Cart/index.tsx + PayModal + ReceiptModal",
    },
    {
        "ma": "D.4.2",
        "title": "Tồn popup biến thể khớp tồn list POS",
        "root_cause": (
            "ProductService.detail không truyền branchId → BE trả tồn tổng "
            "của tất cả kho. Trong khi list POS truyền warehouseId nên tồn "
            "ngoài list là của kho đang chọn → 2 số không nhất quán."
        ),
        "fix": (
            "Truyền branchId (warehouseId) vào useGetDetailProduct → "
            "ProductService.detail → BE filter tồn theo chi nhánh. Popup "
            "biến thể giờ hiển thị tồn của ĐÚNG kho đang chọn."
        ),
        "steps": [
            "Vào Bán hàng tại quầy, chọn kho 'Kho hàng mẫu' (hoặc kho có "
            "SP với biến thể)",
            "Vào danh sách SP → ghi nhớ số tồn hiển thị ngoài card của 1 SP",
            "Click vào SP đó → mở popup biến thể",
            "Verify: tồn trong popup của mỗi biến thể cộng lại phải = tồn "
            "ngoài list (không bị lớn hơn do tính tổng các kho)",
            "Đổi sang kho khác → verify tồn cập nhật đúng theo kho mới",
        ],
        "file": "src/hooks/useGetDetailProduct.ts + ProductGrid + VariantModal + ProductService",
    },
    {
        "ma": "C.1.4",
        "title": "POS chi tiết đơn — không còn hardcode 'Nguyễn Thị Hoa / Bạc'",
        "root_cause": (
            "mappedDataInvoice hardcode tên 'Tên khách hàng', SĐT 'Số điện "
            "thoại', rank 'Bạc'; MOCK_DETAIL_INVOICE có 'Nguyễn Thị Hoa'. "
            "Khi BE chưa trả field → UI hiển thị giá trị giả."
        ),
        "fix": (
            "Bỏ MOCK_DETAIL_INVOICE. mappedDataInvoice đọc từ "
            "invoiceDataApi.customer với nhiều field fallback (name/phone/"
            "points/tier). OrderDetailModal merge thêm customerInfo prop từ "
            "list row làm fallback khi BE chưa trả field."
        ),
        "steps": [
            "Vào Bán hàng & Đơn hàng → Đơn hàng → chọn 1 đơn đã có tên KH",
            "Bấm Xem chi tiết",
            "Verify: tên KH hiển thị ĐÚNG tên thật (không còn 'Nguyễn Thị "
            "Hoa' giả)",
            "SĐT / điểm loyalty / hạng: nếu BE chưa trả về sẽ hiển thị "
            "trống hoặc fallback từ list row (KHÔNG còn 'Số điện thoại' / "
            "'0 điểm' / 'Bạc' fake)",
            "⚠️ Lưu ý: các field SĐT/loyalty/rank cần BE trả về đầy đủ. "
            "Nếu vẫn trống → là do BE chưa cập nhật /invoiceDetail/get.",
        ],
        "file": "src/hooks/useGetDetailInvoice.ts + OrderDetailModal",
    },
    {
        "ma": "C.4.1",
        "title": "Đơn đa kênh lọc theo trạng thái",
        "root_cause": (
            "TAB_STATUS_MAP dùng giá trị sai: 'PROCESSING' (không tồn tại) "
            "thay vì 'SHIPPING', 'REFUNDED' thay vì 'CANCELED'. Filter gửi "
            "giá trị sai lên BE → trả rỗng."
        ),
        "fix": (
            "Đổi 'PROCESSING' → 'SHIPPING', 'REFUNDED' → 'CANCELED' trong "
            "cả TAB_STATUS_MAP và các render case, đồng nhất với status "
            "thực tế BE trả."
        ),
        "steps": [
            "Vào Bán hàng & Đơn hàng → Đơn hàng online → Đơn hàng đa kênh",
            "Click tab 'Chờ xử lý' → list hiển thị đơn có status=PENDING",
            "Click tab 'Đang giao' → list hiển thị đơn có status=SHIPPING "
            "(không còn rỗng)",
            "Click tab 'Hoàn thành' → list hiển thị đơn COMPLETED",
            "Click tab 'Huỷ' → list hiển thị đơn CANCELED (không còn rỗng)",
            "Tab 'Tất cả' → hiển thị mọi đơn bất kể status",
        ],
        "file": "src/pages/MultiChannelSales/MultiChannelOrders/MultiChannelOrders.tsx",
    },
]

for i, bug in enumerate(bugs, 1):
    h = doc.add_heading(f"{i}. {bug['ma']} — {bug['title']}", level=2)

    p = doc.add_paragraph()
    p.add_run("Nguyên nhân: ").bold = True
    p.add_run(bug["root_cause"])

    p = doc.add_paragraph()
    p.add_run("Cách fix: ").bold = True
    p.add_run(bug["fix"])

    p = doc.add_paragraph()
    p.add_run("Bước test:").bold = True
    for step in bug["steps"]:
        doc.add_paragraph(step, style="List Number")

    p = doc.add_paragraph()
    r = p.add_run(f"File: {bug['file']}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

# ── Bug cần chờ BE ──────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading("⏳ Bug cần chờ BE deploy — chưa test được", level=1)

be_bugs = [
    ("C.3.4 / C.3.8", "Chặn tạo 2 phiếu trả hàng cho cùng 1 đơn + cộng tồn",
     "BE cần: (1) chặn server-side khi tạo IV2 trùng referId, (2) cộng tồn "
     "khi confirm IV2. FE đã bỏ defensive re-check cũ (đang block cả legitimate "
     "return)."),
    ("C.3.9", "Số lượng trả hàng luôn hardcode = 1",
     "BE cần trả đúng qty trong /invoice/return/confirm response + "
     "/invoiceDetail/get cho IV2. FE đã populate items[] từ apiItem.products."),
    ("D.4.3", "Trả hàng không cộng kho",
     "BE: xử lý stock adjustment khi confirm IV2."),
    ("D.4.4", "Đổi hàng không cộng/trừ kho",
     "BE: xử lý stock adjustment khi confirm IV11 (cả trả và đổi)."),
    ("C.1.4 (field)", "SĐT, loyalty points, hạng thành viên trên chi tiết đơn",
     "BE cần bổ sung các field customer.phone, customer.points, customer.tier "
     "trong response /invoiceDetail/get. FE đã sẵn sàng hiển thị."),
    ("C.1.5", "Cộng điểm loyalty sau mua",
     "FE đã gọi LoyaltyService.fluctuatePoint sau onPaymentSuccess. Nếu BE "
     "endpoint /market/loyaltyPointLedger/fluctuatePoint hoạt động thì test "
     "được ngay. Công thức: floor(netSpent / exchangeRate)."),
    ("invoice/create 404", "POST /sales/invoice/create 404 khi xác nhận thanh toán",
     "FE đã thêm fallback chain: /create → /create/update → /draft/confirm. "
     "Nếu BE confirm endpoint chính thức thì bỏ fallback."),
]

for ma, title, note in be_bugs:
    p = doc.add_paragraph()
    r = p.add_run(f"{ma} — {title}")
    r.bold = True
    doc.add_paragraph(note, style="Intense Quote")

# ── Footer ──────────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading("Ghi chú retest", level=1)
notes = [
    "Sau khi FE deploy bản build mới (commit 81b76120), vui lòng clear cache "
    "browser (Ctrl+Shift+R) trước khi test.",
    "Các bug C.3.4/C.3.8/C.3.9/D.4.3/D.4.4 đang block lẫn nhau: nếu tạo "
    "phiếu trả hàng thất bại thì các bug cộng tồn / qty liên quan cũng "
    "không test được. Đợi BE fix tạo return trước.",
    "Với F.6.2, tester cần test trên cả desktop và mobile để verify 2 flow "
    "khác nhau (clipboard image vs native share file).",
    "Với E.1.1, workaround client-side sẽ fetch rộng (200 bản ghi) khi "
    "detect keyword là email/phone — nếu tenant có rất nhiều customer thì "
    "performance có thể chậm, đây là trade-off chấp nhận được cho đến khi "
    "BE fix filter.",
    "C.1.5 (loyalty earn) gọi fluctuatePoint với số dương. Nếu BE trả lỗi "
    "thì FE catch và không crash, nhưng tester cần check DB / ledger xem "
    "có row mới không.",
]
for note in notes:
    doc.add_paragraph(note, style="List Bullet")

# Save
output_path = "docs/bugs/retest-guide-2026-04-15.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
