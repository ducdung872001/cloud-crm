"""
Markdown -> DOCX converter for HDSD Reborn CRM.

Designed to produce a professional, customer-ready document similar in quality
to what Claude.ai's web interface generates. Uses python-docx with carefully
crafted styles (no Pandoc).

Usage:
    python md_to_docx.py <input.md> <output.docx>
"""
import os
import re
import sys
from datetime import datetime
from io import BytesIO

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token


# ─── Color palette ─────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x3A, 0x68)        # H1 + cover
TEAL = RGBColor(0x2C, 0x7A, 0x7B)        # H2
DARK = RGBColor(0x2D, 0x3A, 0x4B)        # H3
BODY_GREY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GREY = RGBColor(0xF3, 0xF4, 0xF6)
TABLE_HEADER_BG = "1F3A68"               # navy hex (no #)
TABLE_ALT_BG = "F7F8FA"
CODE_BG = "F3F4F6"
INLINE_CODE_BG = "EFF1F4"
LINK_BLUE = RGBColor(0x1C, 0x5D, 0xB8)
WARN_BG = "FFF7E6"
WARN_BORDER = "F59E0B"


# ─── Helpers ───────────────────────────────────────────────────────────────
def set_cell_background(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="C8CDD3", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_paragraph_shading(paragraph, hex_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def set_paragraph_border(paragraph, color="DDDDDD", left_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    edges = {"top": color, "left": left_color or color, "bottom": color, "right": color}
    sizes = {"top": "4", "left": "24" if left_color else "4", "bottom": "4", "right": "4"}
    for edge, c in edges.items():
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sizes[edge])
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), c)
        pBdr.append(b)
    p_pr.append(pBdr)


def _make_simple_field(field_name, cached="1"):
    """Return a list of OxmlElements implementing begin/instrText/separate/result/end."""
    elements = []
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    elements.append(fc_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    elements.append(instr)
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    elements.append(fc_sep)
    t = OxmlElement("w:t")
    t.text = cached
    elements.append(t)
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    elements.append(fc_end)
    return elements


def add_page_number(paragraph, prefix="Trang "):
    grey = RGBColor(0x6B, 0x72, 0x80)

    run = paragraph.add_run(prefix)
    run.font.size = Pt(9)
    run.font.color.rgb = grey

    # PAGE field — wrap in its own run so font formatting stays attached
    run_page = paragraph.add_run()
    run_page.font.size = Pt(9)
    run_page.font.color.rgb = grey
    for el in _make_simple_field("PAGE", cached="1"):
        run_page._r.append(el)

    sep = paragraph.add_run(" / ")
    sep.font.size = Pt(9)
    sep.font.color.rgb = grey

    run_total = paragraph.add_run()
    run_total.font.size = Pt(9)
    run_total.font.color.rgb = grey
    for el in _make_simple_field("NUMPAGES", cached="1"):
        run_total._r.append(el)


def add_toc(paragraph):
    """Insert a Word TOC field. User must press F9 in Word to populate."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Bấm chuột phải vào đây và chọn 'Update Field' (hoặc F9) để cập nhật mục lục."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def set_run_inline_code(run):
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), INLINE_CODE_BG)
    rPr.append(shd)


# ─── Image sizing ──────────────────────────────────────────────────────────
def get_image_size_for_page(image_path, max_width_inches=6.5):
    """Return (width, height) in EMU sized to fit page width."""
    try:
        with PILImage.open(image_path) as img:
            w_px, h_px = img.size
            dpi = img.info.get("dpi", (96, 96))[0] or 96
        w_in = w_px / dpi
        h_in = h_px / dpi
        if w_in > max_width_inches:
            ratio = max_width_inches / w_in
            w_in = max_width_inches
            h_in = h_in * ratio
        return Inches(w_in), Inches(h_in)
    except Exception:
        return Inches(6), Inches(4)


# ─── Setup styles on a fresh doc ───────────────────────────────────────────
def setup_document_styles(doc: Document):
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_GREY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = TEAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = DARK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ─── Cover page ────────────────────────────────────────────────────────────
def add_cover(doc: Document, title: str, subtitle: str):
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HƯỚNG DẪN SỬ DỤNG")
    r.font.size = Pt(14)
    r.font.color.rgb = TEAL
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(32)
    r.font.color.rgb = NAVY
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = BODY_GREY
    r.font.italic = True

    # Decorative line
    for _ in range(2):
        doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("─" * 40)
    r.font.color.rgb = TEAL
    r.font.size = Pt(11)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu dành cho người dùng cuối")
    r.font.size = Pt(12)
    r.font.color.rgb = BODY_GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Phiên bản: {datetime.now().strftime('%d/%m/%Y')}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Page break
    doc.add_page_break()

    # TOC page
    p = doc.add_paragraph()
    r = p.add_run("MỤC LỤC")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)

    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    doc.add_page_break()


# ─── Inline rendering ──────────────────────────────────────────────────────
def render_inline(paragraph, children, base_dir, list_marker=False):
    """Render markdown_it inline tokens into a python-docx paragraph."""
    if not children:
        return
    bold = False
    italic = False
    code = False
    link = None
    for t in children:
        ttype = t.type
        if ttype == "text":
            text = t.content
            if text:
                run = paragraph.add_run(text)
                if bold: run.bold = True
                if italic: run.italic = True
                if code: set_run_inline_code(run)
                if link is not None:
                    run.font.color.rgb = LINK_BLUE
                    run.font.underline = True
        elif ttype == "strong_open": bold = True
        elif ttype == "strong_close": bold = False
        elif ttype == "em_open": italic = True
        elif ttype == "em_close": italic = False
        elif ttype == "code_inline":
            run = paragraph.add_run(t.content)
            set_run_inline_code(run)
        elif ttype == "link_open":
            link = t.attrs.get("href", "")
        elif ttype == "link_close":
            link = None
        elif ttype == "softbreak":
            paragraph.add_run(" ")
        elif ttype == "hardbreak":
            paragraph.add_run().add_break()
        elif ttype == "image":
            # inline image — render in its own line below current paragraph
            src = t.attrs.get("src", "")
            alt = t.content or "Hình minh họa"
            insert_image(paragraph._parent, src, alt, base_dir)


def insert_image(container, src, alt, base_dir, max_width_inches=6.3):
    """Insert image with caption. container is doc or _Cell."""
    src_path = src
    if not os.path.isabs(src_path):
        src_path = os.path.normpath(os.path.join(base_dir, src.lstrip("./")))
    if not os.path.exists(src_path):
        # Try alternative resolution
        alt_path = os.path.normpath(os.path.join(base_dir, src))
        if os.path.exists(alt_path):
            src_path = alt_path
        else:
            p = container.add_paragraph()
            r = p.add_run(f"[Không tìm thấy ảnh: {src}]")
            r.font.italic = True
            r.font.color.rgb = RGBColor(0xB0, 0x10, 0x10)
            return
    try:
        w, h = get_image_size_for_page(src_path, max_width_inches)
        p = container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(src_path, width=w)
        # caption
        if alt and alt != "Hình minh họa":
            cap = container.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            cr = cap.add_run(f"Hình: {alt}")
            cr.font.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    except Exception as e:
        p = container.add_paragraph()
        r = p.add_run(f"[Lỗi chèn ảnh {src}: {e}]")
        r.font.italic = True
        r.font.color.rgb = RGBColor(0xB0, 0x10, 0x10)


# ─── Token stream parser (block level) ────────────────────────────────────
def render_tokens(tokens, doc, base_dir):
    i = 0
    n = len(tokens)
    while i < n:
        t = tokens[i]
        ttype = t.type

        if ttype == "heading_open":
            level = int(t.tag[1])
            i += 1
            inline = tokens[i]
            text = "".join([c.content for c in inline.children if c.type == "text" or c.type == "code_inline"])
            text = re.sub(r"\s+", " ", text).strip()
            # H1 = page break before (each Part starts on new page)
            if level == 1:
                doc.add_page_break()
            style_name = f"Heading {min(level, 3)}"
            p = doc.add_paragraph(style=style_name)
            # Render inline (to keep bold/code formatting in headings)
            render_inline(p, inline.children, base_dir)
            i += 2  # skip heading_close
            continue

        elif ttype == "paragraph_open":
            i += 1
            inline = tokens[i]
            # Special case: paragraph contains ONLY a single image
            if (len(inline.children) == 1 and inline.children[0].type == "image"):
                img = inline.children[0]
                insert_image(doc, img.attrs.get("src", ""), img.content, base_dir)
            elif (len(inline.children) >= 1 and
                  all(c.type in ("image", "softbreak") for c in inline.children)):
                for c in inline.children:
                    if c.type == "image":
                        insert_image(doc, c.attrs.get("src", ""), c.content, base_dir)
            else:
                p = doc.add_paragraph()
                render_inline(p, inline.children, base_dir)
            i += 2  # skip paragraph_close
            continue

        elif ttype == "bullet_list_open":
            i = render_list(tokens, i, doc, base_dir, ordered=False, level=0)
            continue

        elif ttype == "ordered_list_open":
            i = render_list(tokens, i, doc, base_dir, ordered=True, level=0)
            continue

        elif ttype == "blockquote_open":
            i = render_blockquote(tokens, i, doc, base_dir)
            continue

        elif ttype == "fence" or ttype == "code_block":
            content = t.content.rstrip("\n")
            render_code_block(doc, content)
            i += 1
            continue

        elif ttype == "table_open":
            i = render_table(tokens, i, doc, base_dir)
            continue

        elif ttype == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("─" * 30)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            r.font.size = Pt(10)
            i += 1
            continue

        elif ttype == "html_block":
            # Skip raw HTML
            i += 1
            continue

        else:
            i += 1


def render_list(tokens, i, doc, base_dir, ordered, level):
    """Render bullet/ordered list. Returns next index."""
    open_type = "ordered_list_open" if ordered else "bullet_list_open"
    close_type = "ordered_list_close" if ordered else "bullet_list_close"
    style = "List Number" if ordered else "List Bullet"

    i += 1  # skip the open
    counter = 1
    while i < len(tokens) and tokens[i].type != close_type:
        if tokens[i].type == "list_item_open":
            i += 1
            # The item contains paragraphs / nested lists
            first_para = True
            while i < len(tokens) and tokens[i].type != "list_item_close":
                t = tokens[i]
                if t.type == "paragraph_open":
                    i += 1
                    inline = tokens[i]
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
                    p.paragraph_format.first_line_indent = Cm(-0.4)
                    p.paragraph_format.space_after = Pt(2)
                    if first_para:
                        marker = f"{counter}. " if ordered else "• "
                        mr = p.add_run(marker)
                        mr.font.bold = ordered
                        if not ordered:
                            mr.font.color.rgb = TEAL
                        first_para = False
                    else:
                        p.add_run("   ")
                    render_inline(p, inline.children, base_dir)
                    i += 2  # paragraph_close
                elif t.type == "bullet_list_open":
                    i = render_list(tokens, i, doc, base_dir, ordered=False, level=level + 1)
                elif t.type == "ordered_list_open":
                    i = render_list(tokens, i, doc, base_dir, ordered=True, level=level + 1)
                elif t.type == "fence" or t.type == "code_block":
                    render_code_block(doc, t.content.rstrip("\n"))
                    i += 1
                else:
                    i += 1
            counter += 1
            i += 1  # list_item_close
        else:
            i += 1
    return i + 1  # skip list close


def render_blockquote(tokens, i, doc, base_dir):
    i += 1
    while i < len(tokens) and tokens[i].type != "blockquote_close":
        t = tokens[i]
        if t.type == "paragraph_open":
            i += 1
            inline = tokens[i]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_shading(p, WARN_BG)
            set_paragraph_border(p, color="FCD9A0", left_color=WARN_BORDER)
            render_inline(p, inline.children, base_dir)
            i += 2
        else:
            i += 1
    return i + 1


def render_code_block(doc, content):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_shading(p, CODE_BG)
    set_paragraph_border(p, color="DDDDDD")
    r = p.add_run(content)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4B)


def render_table(tokens, i, doc, base_dir):
    """Render a markdown table."""
    # Collect rows
    rows = []  # list of list of (children inline tokens)
    is_header = True
    header_row = None
    body_rows = []
    i += 1  # table_open
    while i < len(tokens) and tokens[i].type != "table_close":
        t = tokens[i]
        if t.type == "thead_open":
            i += 1
            while tokens[i].type != "thead_close":
                if tokens[i].type == "tr_open":
                    row = []
                    i += 1
                    while tokens[i].type != "tr_close":
                        if tokens[i].type in ("th_open", "td_open"):
                            i += 1
                            inline = tokens[i]
                            row.append(inline.children or [])
                            i += 1  # inline
                            i += 1  # th_close
                        else:
                            i += 1
                    header_row = row
                    i += 1  # tr_close
                else:
                    i += 1
            i += 1  # thead_close
        elif t.type == "tbody_open":
            i += 1
            while tokens[i].type != "tbody_close":
                if tokens[i].type == "tr_open":
                    row = []
                    i += 1
                    while tokens[i].type != "tr_close":
                        if tokens[i].type in ("th_open", "td_open"):
                            i += 1
                            inline = tokens[i]
                            row.append(inline.children or [])
                            i += 1  # inline
                            i += 1  # td_close
                        else:
                            i += 1
                    body_rows.append(row)
                    i += 1  # tr_close
                else:
                    i += 1
            i += 1  # tbody_close
        else:
            i += 1
    # Build docx table
    if not header_row:
        return i + 1
    n_cols = len(header_row)
    table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # Header row
    for ci, cell_children in enumerate(header_row):
        cell = table.rows[0].cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, TABLE_HEADER_BG)
        set_cell_borders(cell)
        # clear default empty paragraph
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        render_inline(p, cell_children, base_dir)
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    # Body rows
    for ri, row in enumerate(body_rows):
        is_alt = (ri % 2 == 1)
        for ci in range(n_cols):
            cell = table.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if is_alt:
                set_cell_background(cell, TABLE_ALT_BG)
            set_cell_borders(cell)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            children = row[ci] if ci < len(row) else []
            render_inline(p, children, base_dir)
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(10)
    # Spacer paragraph after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return i + 1


# ─── Header / footer ───────────────────────────────────────────────────────
def setup_header_footer(doc, doc_title):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header (non-first-page)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(doc_title)
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    hr.font.italic = True

    # Footer (non-first-page)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


# ─── Main ──────────────────────────────────────────────────────────────────
def convert(md_path, docx_path):
    base_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    md = MarkdownIt("commonmark", {"html": True}).enable(["table", "strikethrough"])
    tokens = md.parse(md_text)

    doc = Document()
    setup_document_styles(doc)
    setup_header_footer(doc, "Hướng dẫn sử dụng Reborn CRM")

    add_cover(doc,
              title="Reborn CRM",
              subtitle="Tài liệu hướng dẫn sử dụng — Cửa hàng & Spa")

    # Skip the first H1 (the cover title) and the header content before "Toàn bộ nội dung HDSD"
    # Find the marker and skip everything before
    start_idx = 0
    for idx, t in enumerate(tokens):
        if t.type == "heading_open" and t.tag == "h1":
            # Look at the next inline token
            inline = tokens[idx + 1]
            text = "".join([c.content for c in inline.children if c.type == "text"]).strip()
            if "Part 01" in text or "Bắt đầu sử dụng" in text:
                start_idx = idx
                break

    render_tokens(tokens[start_idx:], doc, base_dir)

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    out = convert(sys.argv[1], sys.argv[2])
    print(f"Done: {out}")
