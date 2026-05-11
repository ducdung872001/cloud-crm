# -*- coding: utf-8 -*-
"""
Build a BoD-grade DOCX from the FitPro 7-chapter documentation.

Input: docs/fitpro/{01-07}-*/{*.md} (concatenated in numeric order) + glossary.md
Output: docs/fitpro/exports/FitPro_2027_Phan_Tich_Tong_Hop.docx

Flow:
  1. Concatenate all chapter markdowns into a single source file.
  2. Customize Pandoc's default reference.docx (fonts, heading colors).
  3. Run Pandoc to convert MD → DOCX using the customized reference.
  4. Post-process: cover page, static TOC, footer with page numbers, table styling.

Run:  python build_docx.py
"""
import os
import sys
import io
import re
import shutil
import subprocess
from glob import glob

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FITPRO_ROOT = os.path.dirname(HERE)
REF_BASE = os.path.join(HERE, "_reference.docx")
REF_OUT = os.path.join(HERE, "_reference_styled.docx")
MD_TMP = os.path.join(HERE, "_source_for_docx.md")
OUT_DOCX = os.path.join(HERE, "FitPro_2027_Phan_Tich_Tong_Hop.docx")

# Brand palette
COLOR_PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
COLOR_ACCENT = RGBColor(0xC9, 0x9A, 0x3B)
COLOR_SUB = RGBColor(0x33, 0x33, 0x33)
COLOR_LIGHT = RGBColor(0x70, 0x70, 0x70)
COLOR_PALE_BG = "F2F4F8"
COLOR_HEADER_BG = "0B3D91"

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"

CHAPTER_TITLES = {
    "01-context":       "Chương 1 — Bối cảnh kinh doanh",
    "02-scope":         "Chương 2 — Phạm vi sản phẩm",
    "03-architecture":  "Chương 3 — Kiến trúc & Mô hình kinh doanh",
    "04-roadmap":       "Chương 4 — Roadmap 2026–2027",
    "05-current-state": "Chương 5 — Hiện trạng & Quyết định mở",
    "06-cost-estimate": "Chương 6 — Sizing · Break-even · Định giá",
    "07-business-case": "Chương 7 — Business Case",
}


def collect_markdown_files():
    files = []
    for chapter in sorted(CHAPTER_TITLES.keys()):
        chapter_dir = os.path.join(FITPRO_ROOT, chapter)
        if not os.path.isdir(chapter_dir):
            continue
        for md in sorted(glob(os.path.join(chapter_dir, "*.md"))):
            files.append((chapter, md))
    gloss = os.path.join(FITPRO_ROOT, "glossary.md")
    if os.path.exists(gloss):
        files.append(("glossary", gloss))
    return files


def concat_markdown():
    files = collect_markdown_files()
    out_lines = []
    current_chapter = None
    for chapter, path in files:
        if chapter != current_chapter:
            current_chapter = chapter
            if chapter == "glossary":
                out_lines.append("\n\n# Glossary\n")
            else:
                title = CHAPTER_TITLES.get(chapter, chapter)
                out_lines.append(f"\n\n# {title}\n")
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        # Demote the first H1 in each file to H2 so it nests under the chapter banner.
        content = re.sub(r"^# ", "## ", content, count=1, flags=re.MULTILINE)
        out_lines.append(content)
        out_lines.append("\n\n---\n\n")
    with open(MD_TMP, "w", encoding="utf-8") as f:
        f.write("".join(out_lines))
    print(f"[ok] concatenated {len(files)} files → {MD_TMP}")


def style_reference():
    shutil.copyfile(REF_BASE, REF_OUT)
    doc = Document(REF_OUT)
    styles = doc.styles

    def set_style(name, *, size=None, bold=None, color=None, font=None, italic=None,
                  space_before=None, space_after=None):
        try:
            st = styles[name]
        except KeyError:
            return
        if font:
            st.font.name = font
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rfonts.set(qn(attr), font)
        if size is not None:
            st.font.size = Pt(size)
        if bold is not None:
            st.font.bold = bold
        if italic is not None:
            st.font.italic = italic
        if color is not None:
            st.font.color.rgb = color
        pfmt = st.paragraph_format
        if space_before is not None:
            pfmt.space_before = Pt(space_before)
        if space_after is not None:
            pfmt.space_after = Pt(space_after)

    set_style("Normal", font=FONT_BODY, size=11, color=COLOR_SUB, space_before=0, space_after=4)
    set_style("Title", font=FONT_HEAD, size=26, bold=True, color=COLOR_PRIMARY, space_after=6)
    set_style("Subtitle", font=FONT_HEAD, size=14, italic=True, color=COLOR_LIGHT, space_after=18)
    set_style("Heading 1", font=FONT_HEAD, size=22, bold=True, color=COLOR_PRIMARY, space_before=24, space_after=10)
    set_style("Heading 2", font=FONT_HEAD, size=16, bold=True, color=COLOR_PRIMARY, space_before=14, space_after=6)
    set_style("Heading 3", font=FONT_HEAD, size=12, bold=True, color=COLOR_ACCENT, space_before=8, space_after=4)
    set_style("Heading 4", font=FONT_HEAD, size=11, bold=True, color=COLOR_SUB, space_before=6, space_after=4)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    doc.save(REF_OUT)
    print(f"[ok] reference styled: {REF_OUT}")


def run_pandoc():
    cmd = [
        "pandoc", MD_TMP,
        "-o", OUT_DOCX,
        "--reference-doc", REF_OUT,
        "-V", "lang=vi",
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        print("STDOUT:", r.stdout)
        print("STDERR:", r.stderr)
        raise SystemExit(r.returncode)
    print(f"[ok] pandoc converted: {OUT_DOCX}")


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def set_cell_border(cell, color="BFBFBF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        for old in tcBorders.findall(qn(f"w:{edge}")):
            tcBorders.remove(old)
        tcBorders.append(el)


def style_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            set_cell_shading(cell, COLOR_HEADER_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, color="0B3D91", size="8")
        for ridx, row in enumerate(table.rows[1:], start=1):
            for cell in row.cells:
                set_cell_border(cell, color="D9DCE3", size="6")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.size is None or run.font.size > Pt(11):
                            run.font.size = Pt(10.5)
                if ridx % 2 == 0:
                    set_cell_shading(cell, COLOR_PALE_BG)


def add_field(paragraph, instr_text):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r_el = run._r
    r_el.append(fld_begin); r_el.append(instr); r_el.append(fld_sep); r_el.append(fld_end)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run("FitPro by MF7 · Vision 2027   |   Trang ")
        run.font.size = Pt(9); run.font.color.rgb = COLOR_LIGHT
        add_field(p, "PAGE")
        run2 = p.add_run(" / ")
        run2.font.size = Pt(9); run2.font.color.rgb = COLOR_LIGHT
        add_field(p, "NUMPAGES")
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_LIGHT


def make_p_oxml(text, *, size=11, bold=False, color="333333", align="center",
                space_before=0, space_after=8, italic=False, font=FONT_HEAD):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), align); pPr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(space_before * 20))
    spacing.set(qn("w:after"), str(space_after * 20))
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    rPr.append(rfonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(size * 2)); rPr.append(szCs)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
        bCs = OxmlElement("w:bCs"); rPr.append(bCs)
    if italic:
        i = OxmlElement("w:i"); rPr.append(i)
    if isinstance(color, RGBColor):
        color_hex = f"{int(color[0]):02X}{int(color[1]):02X}{int(color[2]):02X}"
    else:
        color_hex = color
    col = OxmlElement("w:color"); col.set(qn("w:val"), color_hex); rPr.append(col)
    r.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; r.append(t)
    p.append(r)
    return p


def make_page_break_oxml():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); r.append(br)
    p.append(r)
    return p


def make_hr_oxml():
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C99A3B")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.append(pPr)
    return p


def build_static_toc(doc):
    first_h1 = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            first_h1 = p
            break
    if first_h1 is None:
        return

    entries = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            entries.append((1, p.text.strip()))
        elif p.style.name == "Heading 2":
            entries.append((2, p.text.strip()))

    insert_before = first_h1._element
    parent = insert_before.getparent()

    def add_para(p_el):
        parent.insert(list(parent).index(insert_before), p_el)

    add_para(make_page_break_oxml())

    h = OxmlElement("w:p")
    hpPr = OxmlElement("w:pPr")
    hStyle = OxmlElement("w:pStyle"); hStyle.set(qn("w:val"), "Heading1"); hpPr.append(hStyle)
    h.append(hpPr)
    hr = OxmlElement("w:r")
    hrPr = OxmlElement("w:rPr")
    hb = OxmlElement("w:b"); hrPr.append(hb)
    hsz = OxmlElement("w:sz"); hsz.set(qn("w:val"), "44"); hrPr.append(hsz)
    hcol = OxmlElement("w:color"); hcol.set(qn("w:val"), "0B3D91"); hrPr.append(hcol)
    hr.append(hrPr)
    htext = OxmlElement("w:t"); htext.text = "MỤC LỤC"; hr.append(htext)
    h.append(hr)
    add_para(h)

    for level, text in entries:
        ep = OxmlElement("w:p")
        epPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        if level == 2:
            ind.set(qn("w:left"), "567")
        epPr.append(ind)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        epPr.append(tabs)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "60" if level == 1 else "20")
        spacing.set(qn("w:after"), "40")
        epPr.append(spacing)
        ep.append(epPr)
        er = OxmlElement("w:r")
        erPr = OxmlElement("w:rPr")
        esz = OxmlElement("w:sz"); esz.set(qn("w:val"), "24" if level == 1 else "20"); erPr.append(esz)
        if level == 1:
            eb = OxmlElement("w:b"); erPr.append(eb)
            ecol = OxmlElement("w:color"); ecol.set(qn("w:val"), "0B3D91"); erPr.append(ecol)
        else:
            ecol = OxmlElement("w:color"); ecol.set(qn("w:val"), "333333"); erPr.append(ecol)
        er.append(erPr)
        et = OxmlElement("w:t"); et.set(qn("xml:space"), "preserve"); et.text = text; er.append(et)
        ep.append(er)
        add_para(ep)

    add_para(make_page_break_oxml())
    print(f"[ok] built static TOC with {len(entries)} entries")


def insert_cover_page(doc):
    body = doc.element.body
    elements = [
        make_p_oxml("", space_after=120),
        make_p_oxml("FITPRO BY MF7", size=12, bold=True, color=COLOR_ACCENT, space_after=12),
        make_p_oxml("PHÂN TÍCH TỔNG HỢP CHIẾN LƯỢC", size=28, bold=True, color=COLOR_PRIMARY, space_after=4),
        make_p_oxml("VISION 2027", size=22, bold=True, color=COLOR_PRIMARY, space_after=6),
        make_p_oxml("Kiến tạo 10.000 FitPro Home Toàn Quốc", size=14, italic=True, color=COLOR_LIGHT, space_after=18),
        make_hr_oxml(),
        make_p_oxml("Tài liệu trình Ban Giám đốc", size=13, bold=True, color=COLOR_SUB, space_before=12, space_after=120),
        make_p_oxml("Soạn thảo: CTO Office · Reborn JSC", size=11, color=COLOR_SUB, space_after=4),
        make_p_oxml("Triết lý vận hành: Tốc Độ – Chuẩn Xác – Nhân Văn", size=11, italic=True, color=COLOR_LIGHT, space_after=4),
        make_p_oxml("Ngày 12 tháng 05 năm 2026", size=11, color=COLOR_LIGHT, space_after=4),
    ]
    for el in reversed(elements):
        body.insert(0, el)


def post_process():
    doc = Document(OUT_DOCX)
    build_static_toc(doc)
    insert_cover_page(doc)
    style_tables(doc)
    add_footer(doc)
    doc.save(OUT_DOCX)
    print(f"[ok] post-processed: {OUT_DOCX}")


if __name__ == "__main__":
    concat_markdown()
    style_reference()
    run_pandoc()
    post_process()
    size = os.path.getsize(OUT_DOCX) / 1024
    print(f"[done] {OUT_DOCX} ({size:.1f} KB)")
